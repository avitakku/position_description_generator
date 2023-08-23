from flask import Flask, request, render_template, jsonify, session, send_file

import openai

import pdfplumber
from docx import Document
import io

import config

app = Flask(__name__)

openai.api_key = config.api_key
app.secret_key = config.secret_key

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_files():
    pos_description = ""
    example_pd = ""

    pos_description = request.form.get('position_details')
    example = request.files.get('example')

    example_filename = example.filename
    example_extension = example_filename.rsplit('.', 1)[1].lower()

    if example_extension == 'pdf':
        print('example is pdf\n')

    # Convert the pdf files to text
        with pdfplumber.open(example) as pdf:
            example_pd = '\n'.join(page.extract_text() for page in pdf.pages)
        print(example_pd+'\n')
    
    elif example_extension in ['doc', 'docx']:
        doc = Document(example)

        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text not in paragraphs:
                paragraphs.append(paragraph.text)

        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    #print(cell.text)
                    if cell.text not in tables_text:   
                        tables_text.append(cell.text)

        example_pd = '\n'.join(paragraphs + tables_text)
        print(example_pd+'\n')  # Check the final resume text

    prompt = request.form.get('prompt')
    print(prompt)

    # Make OpenAI API call
    try:
        print('trying model \n')

        response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo-16k",
        messages=[
            {"role": "system", "content": "You are a helpful assistant to a recruiter. Your job is to generate job descriptions \
             based on the position requirements provided to you. Look at the example job description provided to you and write a \
             new job description in the same style. Don't lift content from the example job description-- get all your content from\
             the position requirements. "},
            {"role": "user", "content": f"These are the position requirements based on which you need to generate a new job description:: {pos_description} \
             \n ====== \n Here is an example job description, for inspiration. Don't use content from this: {example_pd}\n ====== \n {prompt}"}
        ]
        )
        
        #Find the last message from the assistant
        assistant_message = response['choices'][0]['message']['content']
        session['assistant_message'] = assistant_message
        
        #print('assistant message: \n')
        #print(assistant_message)
    
    except Exception as e:
        # In case of API error, return the error message
        print(f"Error details: {str(e)}, {type(e).__name__}")
        e = str(e)
        output = "There was an error! Here is the message: " + e
        return jsonify(output)

    return jsonify(assistant_message)

@app.route('/download', methods=['GET'])
def download():
    # Create the .docx file in memory
    f = io.BytesIO()
    doc = Document()
    doc.add_paragraph(session.get('assistant_message', ''))
    doc.save(f)
    f.name = 'results.docx'
    f.seek(0)
    return send_file(f, as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', download_name='results.docx')

if __name__ == "__main__":
    app.run(debug=True)

